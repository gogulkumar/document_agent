"""DOCX parser — extracts paragraphs and table content using python-docx."""
from __future__ import annotations
import logging
logger = logging.getLogger(__name__)

def parse(file_path: str) -> str:
    try:
        from docx import Document  # type: ignore
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX PARSE ERROR: {e}]"

"""task_word_export — generate a Word .docx report using python-docx."""
from __future__ import annotations
import logging
import os
import uuid

from tools.tasks.task_base import invoke_task_llm, EXPORT_ROOT

logger = logging.getLogger(__name__)

_WORD_PROMPT = """
Generate a professional Word document as structured text.
Use these markers for formatting:
  # Heading 1
  ## Heading 2
  ### Heading 3
  - Bullet point
  **Bold text**
  | Col1 | Col2 | (for tables, pipe-separated)

Keep structure clear. Every data point must include its source.
"""

def task_word_export(task_description: str, dependency_payload: str, **kwargs) -> str:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return "ERROR: python-docx not installed"

    raw = invoke_task_llm(task_description + _WORD_PROMPT, dependency_payload, task_type="writing", max_tokens=16000)

    doc = Document()
    doc.add_heading("Research Report", level=0)

    for line in raw.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    path = os.path.join(EXPORT_ROOT, f"{uuid.uuid4().hex}.docx")
    doc.save(path)
    logger.info(f"Word doc saved: {path}")
    return path
